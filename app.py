from fastapi import Form
import json
from fastapi import FastAPI
from pydantic import BaseModel
from rag import (
    build_trunks_and_embeddings,
    analyze_document_structure,
    retrieve_for_summary,
    retrieve_top_k_chunks_with_score,
    embed_query,
    rewrite_query,
    answer_with_rag
)

document = ""
structure = {}
chunks = []
embeddings = []
# embeddings will be built lazily after upload or first query

# ========== 启动时加载默认示例文档 ==========

def load_document(path: str) -> str:
    """
    支持加载 .md / .txt / .docx 作为默认文档
    """
    if path.endswith(".docx"):
        from docx import Document
        import io

        doc = Document(path)
        full_text = []

        for para in doc.paragraphs:
            style = para.style.name
            text = para.text.strip()
            if not text:
                continue

            if style.startswith("Heading 1"):
                full_text.append(f"# {text}")
            elif style.startswith("Heading 2"):
                full_text.append(f"## {text}")
            elif style.startswith("Heading 3"):
                full_text.append(f"### {text}")
            else:
                full_text.append(text)

        return "\n".join(full_text)

    else:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

# 默认示例文档（如果用户没有上传）
document = load_document("doc_example.docx")
structure = analyze_document_structure(document)

# ========== FastAPI ==========

app = FastAPI(title="AI Product Document Assistant")

class AskRequest(BaseModel):
    question: str
    task: str  # summary / risk / advice

@app.post("/ask")
def ask(req: AskRequest):
    global chunks, embeddings, document

    # Lazy build embeddings if not built yet
    if not chunks or len(embeddings) == 0:
        if document:
            print("Building embeddings lazily...")
            chunks, embeddings = build_trunks_and_embeddings(document)
            print("Embeddings built.")

    # 所有任务类型都改写问题
    rewritten_query = rewrite_query(req.question, structure, task=req.task)

    if req.task == "summary":
        results = retrieve_for_summary(
            question=rewritten_query,
            structure=structure,
            chunks=chunks,
            embeddings=embeddings,
            per_module_k=3,
        )
    else:
        query_emb = embed_query(rewritten_query)
        results = retrieve_top_k_chunks_with_score(
            chunks,
            embeddings,
            query_embedding=query_emb,
            top_k=6
        )

    # 统一 results 结构为 (chunk, score, module)
    normalized_results = []
    for item in results:
        if len(item) == 3:
            normalized_results.append(item)
        elif len(item) == 2:
            c, s = item
            normalized_results.append((c, s, "unknown"))

    answer = answer_with_rag(
        question=req.question,
        retrieved_chunks=normalized_results,
        task=req.task
    )

    chunks_for_frontend = [
        {"content": c, "score": s, "module": m} 
        for c, s, m in normalized_results
    ]

    return {
        "rewritten_query": rewritten_query,
        "answer": answer,
        "chunks": chunks_for_frontend
    }

from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse
from fastapi import Request

templates = Jinja2Templates(directory="templates")

@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse(
        "index.html",
        {"request": request}
    )

from fastapi import UploadFile, File
from docx import Document
import io

# 文档解析函数
def parse_md(file_bytes: bytes) -> str:
    """
    解析 Markdown 文件，返回文本（保留标题 # / ## / ###）
    """
    return file_bytes.decode("utf-8")

def parse_docx(file_bytes: bytes) -> str:
    """
    解析 DOCX 文件，返回 Markdown 风格文本（# / ## / ###）
    """
    doc = Document(io.BytesIO(file_bytes))
    full_text = []

    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()
        if not text:
            continue
        # 将 Heading 转成 Markdown
        if style.startswith("Heading 1"):
            full_text.append(f"# {text}")
        elif style.startswith("Heading 2"):
            full_text.append(f"## {text}")
        elif style.startswith("Heading 3"):
            full_text.append(f"### {text}")
        else:
            full_text.append(text)

    return "\n".join(full_text)

# 上传接口
@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    """
    上传文档（支持 .md 和 .docx），并生成 RAG chunk + embedding
    """
    global document, structure, chunks, embeddings

    file_bytes = await file.read()

    if file.filename.endswith(".md"):
        document = parse_md(file_bytes)
    elif file.filename.endswith(".docx"):
        document = parse_docx(file_bytes)
    else:
        return {"error": "Only MD and DOCX files are supported."}

    # 重新分析结构
    structure = analyze_document_structure(document)
    # 生成结构化 chunk + embeddings
    chunks, embeddings = build_trunks_and_embeddings(document)

    return {"message": f"Document '{file.filename}' uploaded and indexed successfully."}


# ========== 新增 /result 路由用于处理表单提交并渲染 result.html ==========
from fastapi import Form
import json

@app.post("/result", response_class=HTMLResponse)
async def result(
    request: Request,
    rewritten_query: str = Form(...),
    answer: str = Form(...),
    chunks: str = Form(...),
    document_content: str = Form(...)
):
    # 尝试解析 chunks 和 answer 字符串为对象
    try:
        chunks_obj = json.loads(chunks)
    except Exception:
        chunks_obj = chunks
    try:
        answer_obj = json.loads(answer)
    except Exception:
        answer_obj = answer
    return templates.TemplateResponse(
        "result.html",
        {
            "request": request,
            "rewritten_query": rewritten_query,
            "answer": answer_obj,
            "chunks": chunks_obj,
            "document_content": document
        }
    )